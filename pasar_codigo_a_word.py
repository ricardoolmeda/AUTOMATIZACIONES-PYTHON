import os
from docx import Document
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Nombre del archivo Word
file_name = "codigo_python.docx"

# 🔹 Si el archivo ya existe, lo eliminamos
if os.path.exists(file_name):
    os.remove(file_name)
    print(f"🗑 Archivo '{file_name}' eliminado.")

# Código en Python que queremos escribir en Word
codigo_python = """import cv2
import mediapipe as mp
import numpy as np

# Inicializar Mediapipe Hands
mp_hands = mp.solutions.hands
hands = mp_hands.Hands(static_image_mode=True, max_num_hands=2, min_detection_confidence=0.3)

mp_drawing = mp.solutions.drawing_utils  # Para dibujar las manos

# Cargar la imagen
image_path = "D:\\Documentos Richard\\Estudios\\MASTER\\TFM\\fotos_pruebas\\photo.jpg"
image = cv2.imread(image_path)

if image is None:
    print("⚠ Error: No se pudo cargar la imagen. Verifica la ruta.")
    exit()
"""

# Crear un nuevo documento de Word
doc = Document()
doc.add_heading('Código en Python', level=1)  # Agregar título

# Agregar el código con formato de texto monoespaciado
code_paragraph = doc.add_paragraph()
run = code_paragraph.add_run(codigo_python)
run.font.name = 'Courier New'  # Fuente monoespaciada
run.font.size = Pt(10)  # Tamaño de fuente

# Aplicar fondo gris al código
shading_elm = parse_xml(r'<w:shd {} w:fill="D9D9D9"/>'.format(nsdecls('w')))
code_paragraph._element.get_or_add_pPr().append(shading_elm)

# Guardar el documento
doc.save(file_name)
print(f"✅ Documento Word creado: '{file_name}'")
